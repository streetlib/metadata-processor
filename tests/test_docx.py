# -*- coding: UTF-8 -*-

# self.assertEqual(a, b) – check a and b are equal
# self.assertNotEqual(a, b) – check a and b are not equal
# self.assertIn(a, b) – check that a is in the item b
# self.assertNotIn(a, b) – check that a is not in the item b
# self.assertFalse(a) – check that the value of a is False
# self.assertTrue(a) – check the value of a is True
# self.assertIsInstance(a, TYPE) – check that a is of type “TYPE”
# self.assertRaises(ERROR, a, args) – check that when a is called with args that it raises ERROR


import	unittest
import	os
from	pprint		import	pprint

from	docx		import	fullyQualifiedName
from	docx		import	DocxProcessor
from	document	import	Document
from	fragments	import	*

class DocxTests (unittest.TestCase):

	def fragment(self, lineNumber, range, text, attributes):
		if range:
			result = LineChunkFragmentIdentifier().initWithDetails(FragmentIdentifier.getNextIdentifier(), lineNumber, range, text, attributes)
		else:
			result = WholeLineFragmentIdentifier().initWithDetails(FragmentIdentifier.getNextIdentifier(), (lineNumber, lineNumber), text, attributes)

		return result


	def test_fullyQualifiedName (self):
		self.assertEqual(fullyQualifiedName('w', 'pPr'), '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')


	def test_paragraphStyleReading (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '02_paragraph_styles.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		document = docxProcessor._word_document
		paragraph = [element for element in document.iter() if element.tag == fullyQualifiedName('w', 'p')][0]
		pStyle = paragraph.find('./' + fullyQualifiedName('w','pPr') + '/' + fullyQualifiedName('w', 'pStyle'))
		style = pStyle.get(fullyQualifiedName('w', 'val'))
		self.assertEqual(style, 'style1')


	def test_paragraphStyleMapping (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '02_paragraph_styles.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		styleMapping = docxProcessor.paragraphStyleMapping()
		self.assertEqual(len(styleMapping), 9)
		self.assertIn('style0', styleMapping.keys())
		self.assertEqual('Default Style', styleMapping['style0'])
		self.assertIn('style1', styleMapping.keys())
		self.assertEqual('Heading 1', styleMapping['style1'])


	def test_formattingFragmentSorting (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '03_text_formatting.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		document = docxProcessor.document()

		expectedFormatting = [
			self.fragment(1, None, "I can do that\n", {'style': 'Heading 1'}),
			self.fragment(3, None, "I gotta piss\n", {'style': 'Heading 1'}),
			self.fragment(4, (133, 158), "Do you believe that shit?", {'formattings': ['underline']}),
			self.fragment(5, None, "We happy?\n", {'style': 'Heading 1'}),
			self.fragment(6, (118, 269), "Do you see a little Asian child with a blank expression on his face sitting outside on a mechanical helicopter that shakes when you put quarters in it?", {'formattings': ['bold']}),
			self.fragment(7, None, "Uuummmm, this is a tasty burger!\n", {'style': 'Heading 1'}),
			self.fragment(8, (335, 347), "motherfucker", {'formattings': ['italic']}),
			self.fragment(8, (443, 455), "motherfucker", {'formattings': ['italic']}),
			self.fragment(9, None, "I'm serious as a heart attack\n", {'style': 'Heading 1'}),
			self.fragment(10, (182, 194), "Motherfucker", {'formattings': ['italic']})
		]

		print("DOCUMENT FORMATTING:\n" + str(document.formatting()))
		self.assertIsInstance(document.formatting()[0], FragmentIdentifier)
		self.assertEqual(expectedFormatting, document.formatting())


	def test_textTrimming (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '07_trim.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		document = docxProcessor.document()
		
		self.assertEqual('Paragraph with trailing and closing spaces', document.content()[0])
		self.assertEqual('Paragraph with trailing and closing spaces', document.formatting()[0].text),


	def test_textWithLineBreak (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '06_new_line_same_paragraph.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		document = docxProcessor.document()

		self.assertEqual('Title', document.content()[0])
		self.assertEqual('Regular paragraph.\rNew line on the same paragraph.', document.content()[1])
		self.assertEqual('New paragraph.', document.content()[2])


	def test_textWithLineBlocks (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '08_line_blocks.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		document = docxProcessor.document()
		expectedDocument = Document().initWithFile(os.path.join(os.getcwd(), 'samples', 'expected outcome', 'docx', 'test_08'))

		self.assertEquals(expectedDocument.content(),	 document.content())
		self.assertEquals(expectedDocument.formatting(), document.formatting())


	def test_toc (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '09_toc.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		document = docxProcessor.document()
		expectedDocument = Document().initWithFile(os.path.join(os.getcwd(), 'samples', 'expected outcome', 'docx', 'test_09'))

		self.assertEquals(expectedDocument.content(),	 document.content())
		self.assertEquals(expectedDocument.formatting(), document.formatting())


	def test_textWithLineBreaks (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '10_line-break.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		document = docxProcessor.document()
		expectedDocument = Document().initWithFile(os.path.join(os.getcwd(), 'samples', 'expected outcome', 'docx', 'test_10'))

		self.assertEquals(expectedDocument.content(),	 document.content())
		self.assertEquals(expectedDocument.formatting(), document.formatting())


	def test_textWithWeirdFormatting (self):
		sampleDocxFile = os.path.join(os.getcwd(), 'samples', 'docx', '11_weird_formatting.docx')
		docxProcessor = DocxProcessor(sampleDocxFile)
		document = docxProcessor.document()
		expectedDocument = Document().initWithFile(os.path.join(os.getcwd(), 'samples', 'expected outcome', 'docx', 'test_11'))

		self.assertEquals(expectedDocument.content(), document.content())
		self.assertEquals(expectedDocument.formatting(), document.formatting())

